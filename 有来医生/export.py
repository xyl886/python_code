# -*- coding: utf-8 -*-
import os

from docx import Document

from mongodb_tool import MongoDB

os.makedirs('data', exist_ok=True)
db = MongoDB(db_name='youlai')
question_c = db['question']
record_c = db['record']
# 查询数据库中的所有文档
documents = question_c.find_documents()

# 创建一个新的Document对象
doc = Document()

# 将数据写入docx
for index, doc_data in enumerate(documents, start=1):
    doc.add_paragraph(f"{index}. {doc_data["title"]}")  # 序号 + 问题
    doc.add_paragraph(doc_data["description"])
    doc.add_paragraph("病情回答：" + doc_data["answer"])

# 保存文件
doc.save(f"data/youlai.docx")
