# -*- coding: utf-8 -*-

import requests
from docx import Document

from mongodb_tool import MongoDB

headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36 MicroMessenger/7.0.20.1781(0x6700143B) NetType/WIFI MiniProgramEnv/Windows WindowsWechat/WMPF WindowsWechat(0x63090a13)XWEB/11253',
    'xweb_xhr': '1',
    'Authorization': 'eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJ1c2VybmFtZSI6ImhvdF8gbW91c2UiLCJzaWQiOjYyNjU0LCJ1aWQiOjIwMjIzMTcsIm9wZW5pZCI6Im9rRXlKNVphQVRiRUJXODJsM2dJQ3E1aTNWVVkiLCJhcHBpZCI6Ind4NjQwYTdkMTdhZDk2ZTQwNyIsImV4cCI6MTcyNjMyMTMzMCwiaWF0IjoxNzI1NDU3MzMwLCJpc3MiOiJmdWNrZXIifQ.MDW8AMsLPGcMOs5h78F_LFPeEFBriBrWfRSTH5sg6gY',
    'Content-Type': 'application/json',
    'Sec-Fetch-Site': 'cross-site',
    'Sec-Fetch-Mode': 'cors',
    'Sec-Fetch-Dest': 'empty',
    'Referer': 'https://servicewechat.com/wx640a7d17ad96e407/28/page-frame.html',
    'Accept-Language': 'zh-CN,zh;q=0.9',
}


def showsub():
    params = {
        'appid': 'wx640a7d17ad96e407',
    }
    response = requests.get('https://as.kyexam.com/showsub', params=params, headers=headers)
    res_json = response.json()
    showsub_c.save_dict_list_to_collection(res_json['data'], 'id')


def main(bid):
    params = {
        'start': '0',
        'limit': '100',
        'bid': bid,
        'type': '0',
    }

    response = requests.get('https://as.kyexam.com/api/v1/special/lists', params=params, headers=headers)
    print(response.json())
    list_c.save_dict_list_to_collection(response.json()['data'], 'id')


def curl(id):
    params = {
        'bid': id,
        'needpay': '0',
        'reset': '0',
    }

    response = requests.get('https://as.kyexam.com/api/v1/special/start', params=params, headers=headers)
    if not response.json()['data']:
        return
    print(len(response.json()['data']['allqid']))
    ids = [i['qid'] for i in response.json()['data']['allqid']]
    get_answer(id, ids)


def get_answer(bid, ids):
    data = {
        'qids': ",".join(map(str, ids)),
    }
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36 MicroMessenger/7.0.20.1781(0x6700143B) NetType/WIFI MiniProgramEnv/Windows WindowsWechat/WMPF WindowsWechat(0x63090a13)XWEB/11253',
        'Content-Type': 'application/x-www-form-urlencoded',
        'xweb_xhr': '1',
        'Authorization': 'eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJ1c2VybmFtZSI6ImhvdF8gbW91c2UiLCJzaWQiOjYyNjU0LCJ1aWQiOjIwMjIzMTcsIm9wZW5pZCI6Im9rRXlKNVphQVRiRUJXODJsM2dJQ3E1aTNWVVkiLCJhcHBpZCI6Ind4NjQwYTdkMTdhZDk2ZTQwNyIsImV4cCI6MTcyNjMyMTMzMCwiaWF0IjoxNzI1NDU3MzMwLCJpc3MiOiJmdWNrZXIifQ.MDW8AMsLPGcMOs5h78F_LFPeEFBriBrWfRSTH5sg6gY',
        'Sec-Fetch-Site': 'cross-site',
        'Sec-Fetch-Mode': 'cors',
        'Sec-Fetch-Dest': 'empty',
        'Referer': 'https://servicewechat.com/wx640a7d17ad96e407/28/page-frame.html',
        'Accept-Language': 'zh-CN,zh;q=0.9',
    }
    response = requests.post('https://as.kyexam.com/api/v1/questions/getquestionsviaqid', headers=headers, data=data)
    data = response.json()['data']
    print(len(data))
    for item in data:
        item['bid'] = bid
    question_c.save_dict_list_to_collection(data, 'qid')


def export():
    # 创建一个新的 docx 文档
    doc = Document()

    # 遍历 list_c 集合
    for test in list_c.find_documents():
        # 获取对应的题目信息
        infos = question_c.find_documents({'bid': test['id']})
        doc.add_heading(f"{test['title']}", level=1)
        index = 0
        # 将练习名和题目信息组织在一起
        for info in infos:
            # 提取题目信息
            title = info.get('title', '未知题目').replace('&nbsp;', ' ')
            feedback = info.get('feedback', '无反馈')
            options = info.get('options', [])
            rightoids = info.get('rightoids', [])

            # 将题目信息写入到 docx 文件
            doc.add_heading(str(index + 1) + "." + title, level=2)

            doc.add_heading("选项:", level=3)
            for idx, option in enumerate(options):
                option_letter = chr(65 + idx)  # 65 是字母 'A' 的 ASCII 码
                doc.add_paragraph(
                    f"{option_letter}. {option['title'].replace('[/b]', '').replace('[b]', '').replace('[sub]', '').replace('[/sub]', '')}")

            doc.add_heading("正确答案:", level=3)
            if rightoids:
                for oid in rightoids:
                    correct_option = next((chr(65 + idx) for idx, opt in enumerate(options) if opt['oid'] == oid), None)
                    if correct_option:
                        doc.add_paragraph(f"正确答案: {correct_option}")

                doc.add_heading("反馈:", level=3)
                doc.add_paragraph(feedback)
            else:
                print(rightoids)
    # 保存 docx 文件
    doc.save('output.docx')


if __name__ == '__main__':
    db = MongoDB(db_name='kyexam')
    showsub_c = db['showsub']
    list_c = db['list']
    question_c = db['question']
    # showsub()
    # for item in showsub_c.find_documents():
    #     main(item['id'])
    for item in list_c.find_documents(query={'question_fetch': {'$exists': False}}):
        curl(item['id'])
        item['question_fetch'] = True
        list_c.save_dict_to_collection(item, 'id')
